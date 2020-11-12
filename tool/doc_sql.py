import pandas as pd
import io
import csv
from docx import Document
import re
import os
import argparse
import cx_Oracle


def read_docx_tables(filename, tab_id=None, **kwargs):
    """
    parse table(s) from a Word Document (.docx) into Pandas DataFrame(s)

    Parameters:
        filename:   file name of a Word Document

        tab_id:     parse a single table with the index: [tab_id] (counting from 0).
                    When [None] - return a list of DataFrames (parse all tables)

        kwargs:     arguments to pass to `pd.read_csv()` function

    Return: a single DataFrame if tab_id != None or a list of DataFrames otherwise
    """

    def read_docx_tab(tab, **kwargs):
        vf = io.StringIO()
        writer = csv.writer(vf)
        for row in tab.rows:
            writer.writerow(cell.text for cell in row.cells)
        vf.seek(0)
        return pd.read_csv(vf, **kwargs)

    doc = Document(filename)

    if tab_id is None:
        return [read_docx_tab(tab, **kwargs) for tab in doc.tables]
    else:
        try:
            return read_docx_tab(doc.tables[tab_id], **kwargs)
        except IndexError:
            print('Error: specified [tab_id]: {}  does not exist.'.format(tab_id))
            raise


def extract_table_name(basename):
    return "_".join(basename.split("_")[:-1])


def extract_sql(df, basename, tablespace):
    table_comment = basename.split("_")[-1].split('.')[0]
    table_name = "_".join(basename.split("_")[:-1])

    df = df.fillna('')
    filtered_df = df.loc[df['字段名'] != df['定义']]
    if u"默认值" in filtered_df:
        default_values = filtered_df[u"默认值"].map(
            lambda v: 'default ' + v.replace('‘', "'").replace('’', "'") if v else '')
        table_sql = "  " + filtered_df[u"字段名"] + " " + filtered_df[u"定义"] + " " + default_values + " " + filtered_df[
            u"空/非空"]
    else:
        table_sql = "  " + filtered_df[u"字段名"] + " " + filtered_df[u"定义"] + " " + filtered_df[u"空/非空"]
    table_sql = table_sql.str.cat(sep=',\n')
    table_sql = "create table {} \n(\n{} \n) \n" \
                "tablespace {} \n" \
                "  pctfree 10 \n" \
                "  initrans 1 \n" \
                "  maxtrans 255 \n" \
                "  storage \n" \
                "  ( \n" \
                "    initial 64K \n" \
                "    minextents 1 \n" \
                "    maxextents unlimited \n" \
                "  ) " \
                ";\n\n".format(table_name, table_sql, tablespace)

    tb_comment = "comment on column " + table_name + "." + filtered_df[u"字段名"] + " is '" + filtered_df[u"说明"] + "';"
    tb_comment = tb_comment.str.cat(sep='\n')
    tb_comment = u"comment on table {} is '{}';\n{}\n".format(table_name, table_comment, tb_comment)

    # print(table_sql.replace(u'\xa0', u' '))
    # print(tb_comment.replace(u'\xa0', u' '))

    return table_sql + tb_comment


def filter_index(path, basename, index_tablespace):
    doc = Document(path)
    table_name = extract_table_name(basename)

    rt = [parse_index(p.text, 'PK', table_name, index_tablespace) for p in doc.paragraphs] + \
         [parse_index(p.text, 'UK', table_name, index_tablespace) for p in doc.paragraphs] + \
         [parse_index(p.text, 'NK', table_name, index_tablespace) for p in doc.paragraphs]

    rt = '\n'.join(list(filter(None, rt)))
    # print(rt)

    return rt


def get_key_type(name):
    if name.startswith('PK'):
        return "primary key"
    if name.startswith('UK'):
        return "unique"


def parse_index(line, prefix, table_name, index_tablespace):
    line = line.replace('+', ',')
    if line.strip().startswith(prefix):
        m = re.search('([a-zA-Z_ 0-9]*)(\\(.*\\))', line)
        if len(m.groups()) != 2:
            return ''

        key_name = m.group(1).replace(' ', '')
        key_values = m.group(2).replace('+', ',')

        table_space = "tablespace {}\n " \
                      "  pctfree 10\n " \
                      "  initrans 2\n " \
                      "  maxtrans 255\n " \
                      "  storage \n" \
                      "  ( \n" \
                      "    initial 64K \n" \
                      "    minextents 1 \n" \
                      "    maxextents unlimited \n" \
                      "  ) " \
                      ";".format(index_tablespace)

        if prefix == 'NK':
            return "create index {} on {} {}\n".format(key_name, table_name, key_values) + table_space

        return "alter table {} add constraint {} {} {}\n" \
               "using index\n".format(table_name, key_name, get_key_type(key_name),
                                      key_values) + table_space


def extract_all(f, tablespace, index_tablespace, opath, db_cursor):
    basename = os.path.basename(f)
    table_name = "_".join(basename.split("_")[:-1])

    df = read_docx_tables(f, 1)
    table_sql = extract_sql(df, basename, tablespace)
    index_sql = filter_index(f, basename, index_tablespace)

    wf = open("{}/{}.sql".format(opath, basename.replace('.docx', '')), "w+", encoding='GBK')
    wf.writelines(table_sql)
    wf.writelines(index_sql)
    wf.close()

    if db_cursor:
        try:
            db_cursor.execute('drop table {}'.format(table_name))
        except cx_Oracle.DatabaseError as er:
            pass

        for s in table_sql.split(';'):
            try:
                if s.strip():
                    db_cursor.execute(s)
            except Exception as er:
                print('执行失败：{}'.format(s))
                print(er)

        for s in index_sql.split(';'):
            try:
                if s.strip():
                    db_cursor.execute(s)
            except Exception as er:
                print('执行失败：{}'.format(s))
                print(er)


if __name__ == "__main__":
    parser = argparse.ArgumentParser("将表结构word文档生成sql语句")
    parser.add_argument("--tablespace", type=str, help='表空间', default='TBS_CSP_CLI')
    parser.add_argument("--index_tablespace", type=str, help='索引表空间', default='TBS_CSPIDX_DEF')
    parser.add_argument('--input', type=str, help='文档目录', default='.')
    parser.add_argument('--output', type=str, help='输出目录')
    parser.add_argument('--user', type=str, help='数据库用户名')
    parser.add_argument('--password', type=str, help='数据库密码')
    parser.add_argument('--url', type=str, help='数据库地址')

    args = parser.parse_args()

    cursor = None
    if args.user and args.password and args.url:
        connection = cx_Oracle.connect(args.user, args.password, args.url)
        connection.autocommit = True
        cursor = connection.cursor()

    cur_dir = args.input
    opath = args.output or os.path.join(cur_dir,'build')

    if not os.path.exists(opath):
        os.mkdir(opath)

    for f in os.listdir(cur_dir):
        # f = f.decode("gbk")
        try:
            if not f.endswith("docx") or f.startswith("~$"):
                continue

            extract_all(os.path.join(cur_dir, f), args.tablespace, args.index_tablespace, opath, cursor)
        except Exception as e:
            print(e)
